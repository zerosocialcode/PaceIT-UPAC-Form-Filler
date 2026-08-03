import re
import os
import io
import zipfile
import unicodedata
from datetime import datetime

from flask import (
    Flask, request, render_template, redirect, url_for, send_file, flash
)
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)
app.secret_key = "change-me-to-something-random"

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(BASE_DIR, "EDC Form.docx")
OUTPUT_DIR = "/tmp/generated_forms" if os.environ.get("VERCEL") else os.path.join(BASE_DIR, "generated_forms")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def nfc(s):
    return unicodedata.normalize('NFC', s) if s else s


def sanitize_text(s):
    if not s:
        return s
    return s.encode('utf-8', 'ignore').decode('utf-8')


BANGLA_FONT = "Nikosh"
ENGLISH_FONT = "Times New Roman"
DEFAULT_SIZE = Pt(11)


def _is_bangla_char(ch):
    return '\u0980' <= ch <= '\u09FF'


def _char_script(ch):
    if _is_bangla_char(ch):
        return True
    if ch.isalpha() or ch.isdigit():
        return False
    return None


def _split_by_script(text):
    segments = []
    current = ''
    current_is_bangla = None
    for ch in text:
        is_bangla = _char_script(ch)
        if is_bangla is None:
            is_bangla = current_is_bangla if current_is_bangla is not None else False
        if current_is_bangla is None:
            current_is_bangla = is_bangla
            current = ch
        elif is_bangla == current_is_bangla:
            current += ch
        else:
            segments.append((current, current_is_bangla))
            current = ch
            current_is_bangla = is_bangla
    if current:
        segments.append((current, current_is_bangla))
    return segments if segments else [(text, False)]


def _set_run_font(run, font_name):
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), font_name)


def _set_run_size(run, size):
    run.font.size = size
    rPr = run._element.get_or_add_rPr()
    half_points = str(int(round(size.pt * 2)))
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), half_points)


def _copy_run_format(src_run, dst_run):
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    _set_run_size(dst_run, src_run.font.size or DEFAULT_SIZE)
    try:
        if src_run.font.color and src_run.font.color.rgb is not None:
            dst_run.font.color.rgb = src_run.font.color.rgb
    except Exception:
        pass


def apply_mixed_fonts(paragraph, run):
    text = run.text
    if not text:
        return
    segments = _split_by_script(text)

    first_text, first_is_bangla = segments[0]
    run.text = first_text
    _set_run_font(run, BANGLA_FONT if first_is_bangla else ENGLISH_FONT)

    prev_element = run._element
    for seg_text, seg_is_bangla in segments[1:]:
        new_run = paragraph.add_run(seg_text)
        _copy_run_format(run, new_run)
        _set_run_font(new_run, BANGLA_FONT if seg_is_bangla else ENGLISH_FONT)
        prev_element.addnext(new_run._element)
        prev_element = new_run._element


def _capture_format(paragraph):
    if paragraph.runs:
        r = paragraph.runs[0]
        color = None
        try:
            if r.font.color and r.font.color.type is not None:
                color = r.font.color.rgb
        except Exception:
            pass
        return {
            'size': r.font.size or DEFAULT_SIZE,
            'bold': r.font.bold,
            'italic': r.font.italic,
            'underline': r.font.underline,
            'color': color,
        }
    return {'size': DEFAULT_SIZE, 'bold': None, 'italic': None, 'underline': None, 'color': None}


def _apply_captured_format(run, fmt):
    _set_run_size(run, fmt['size'])
    run.font.bold = fmt['bold']
    run.font.italic = fmt['italic']
    run.font.underline = fmt['underline']
    if fmt['color'] is not None:
        try:
            run.font.color.rgb = fmt['color']
        except Exception:
            pass


def set_paragraph_text_preserving_format(paragraph, new_text):
    new_text = sanitize_text(new_text)
    fmt = _capture_format(paragraph)
    paragraph.text = new_text
    if not paragraph.runs:
        return
    run = paragraph.runs[0]
    _apply_captured_format(run, fmt)
    apply_mixed_fonts(paragraph, run)


def to_bangla_digits(number):
    bangla_digits = {'0': '০', '1': '১', '2': '২', '3': '৩', '4': '৪',
                      '5': '৫', '6': '৬', '7': '৭', '8': '৮', '9': '৯'}
    return ''.join(bangla_digits.get(ch, ch) for ch in str(number))


def get_bangla_date(dt):
    day = to_bangla_digits(f"{dt.day:02d}")
    month = to_bangla_digits(f"{dt.month:02d}")
    year = to_bangla_digits(dt.year)
    return f"{day}/{month}/{year}"


def get_today_bangla_date():
    return get_bangla_date(datetime.now())


_MAC_CANDIDATE_RE = re.compile(
    r'(?<![0-9A-Fa-f:\-])'
    r'([0-9A-Fa-f](?:[:\-]?[0-9A-Fa-f]){11})'
    r'(?![0-9A-Fa-f])'
)


def _format_mac(hex_only):
    pairs = [hex_only[i:i + 2] for i in range(0, 12, 2)]
    return ':'.join(pairs).upper()


def normalize_mac(full_text, router_line=""):
    label_match = re.search(r'MAC\s*[:।\-]*\s*([^\n]+)', full_text, re.IGNORECASE)
    if label_match:
        candidate = _MAC_CANDIDATE_RE.search(label_match.group(1))
        if candidate:
            hex_only = re.sub(r'[^0-9A-Fa-f]', '', candidate.group(1))
            if len(hex_only) == 12:
                return _format_mac(hex_only)

    if router_line:
        candidate = _MAC_CANDIDATE_RE.search(router_line)
        if candidate:
            hex_only = re.sub(r'[^0-9A-Fa-f]', '', candidate.group(1))
            if len(hex_only) == 12:
                return _format_mac(hex_only)

    candidate = _MAC_CANDIDATE_RE.search(full_text)
    if candidate:
        hex_only = re.sub(r'[^0-9A-Fa-f]', '', candidate.group(1))
        if len(hex_only) == 12:
            return _format_mac(hex_only)

    return ""


LABELS = {
    "edc_book": ["EDC Book No.:"],
    "connection_nms": ["সংযোগ/NMS আইডিঃ"],
    "address": ["সংযোগ গ্রহণকারী প্রতিষ্টানের নাম ও ঠিকানাঃ"],
    "lat_long": ["Lat & Long:"],
    "union": ["ইউনিয়ন/পৌরসভা (ওয়ার্ড):"],
    "upazila": ["উপজেলা/সিটি কর্পোরেশনঃ"],
    "institution_code": ["প্রতিষ্টানের কোড (EMIS/EIIN/OTHERS):"],
    "date": ["সংযোগ প্রদানের তারিখঃ"],
    "router_mac": ["Wi-Fi Router MAC Address(Upon which API will be called):"],
    "router_serial": ["Wi-Fi Router Serial No.:"],
    "cable_brand": ["Optical Fiber Cable Brand:"]
}

HEADER_SOURCE_NO_LABELS = ["সুত্র নংঃ", "সুত্র নং:", "সুত্র নং"]
HEADER_DATE_LABELS = ["তারিখঃ"]

_RECORD_SPLIT_RE = re.compile(r'^\s*another\s*$', re.IGNORECASE | re.MULTILINE)


def split_records(raw_text):
    blocks = _RECORD_SPLIT_RE.split(raw_text)
    return [b.strip('\n') for b in blocks if b.strip()]


def parse_raw_data(text):
    data = {}
    text = nfc(text)

    def find_field(pattern, default=""):
        match = re.search(nfc(pattern), text, re.IGNORECASE | re.DOTALL)
        return match.group(1).strip() if match else default

    data['serial'] = find_field(r'প্রতিষ্ঠানের\s*সিরিয়াল\s*নাম্বার\s*[:।\-]*\s*([^\n]+)')
    data['name'] = find_field(r'প্রতিষ্ঠানের\s*নাম\s*[:।\-]*\s*([^\n]+)')
    data['mobile'] = find_field(r'মোবাইল\s*নাম্বার\s*[:।\-]*\s*([^\n]+)')
    data['village'] = find_field(r'গ্রাম\s*[:।\-]*\s*([^\n]+)')
    data['pop'] = find_field(r'পপ\s*এর\s*ঠিকানা\s*[:।\-]*\s*([^\n]+)')
    data['union'] = find_field(r'ইউনিয়নের\s*নাম\s*[:।\-]*\s*([^\n]+)')
    data['upazila'] = find_field(r'থানা\s*[:।\-]*\s*([^\n]+)')

    router_line = find_field(r'রাউটার\s*সিরিয়াল\s*নাম্বার\s*[:।\-]*\s*([^\n]+)')
    serial = ""
    if router_line:
        serial = re.sub(nfc(r'^(S/N|SN|সিরিয়াল)\s*[:।\-]?\s*'), '', router_line, flags=re.IGNORECASE).strip()

    mac = normalize_mac(text, router_line)

    if not serial and router_line:
        if mac:
            stripped = re.sub(r'[0-9A-Fa-f]{2}[:\-]?[0-9A-Fa-f]{2}[:\-]?[0-9A-Fa-f]{2}[:\-]?[0-9A-Fa-f]{2}[:\-]?[0-9A-Fa-f]{2}[:\-]?[0-9A-Fa-f]{2}', '', router_line)
            serial = re.sub(nfc(r'^(S/N|SN|সিরিয়াল)\s*[:।\-]?\s*'), '', stripped, flags=re.IGNORECASE).strip()
        else:
            serial = router_line.strip()

    data['router_serial'] = serial
    data['router_mac'] = mac

    return data


FIELD_PROMPTS = [
    ('name', "প্রতিষ্ঠানের নাম"),
    ('mobile', "মোবাইল নাম্বার"),
    ('village', "গ্রাম"),
    ('pop', "পপ এর ঠিকানা"),
    ('union', "ইউনিয়নের নাম"),
    ('upazila', "থানা/উপজেলা"),
    ('router_serial', "রাউটার সিরিয়াল নাম্বার"),
    ('router_mac', "MAC address"),
]

ALWAYS_MANUAL_FIELDS = [
    ('source_no', "EDC Book No. / সুত্র নং", "e.g., 2026-001", ""),
    ('connection_nms', "সংযোগ/NMS আইডিঃ", "", ""),
    ('institution_code', "প্রতিষ্টানের কোড (EMIS/EIIN/OTHERS)", "", ""),
    ('cable_brand', "Optical Fiber Cable Brand", "default 'BRB'", "BRB"),
    ('cable_qty', "Optical Fiber Cable Qty. in meters", "default '1500'", "1500"),
    ('lat_long', "Lat & Long (if available)", "", ""),
]


def _find_label(text, label_variants):
    ntext = nfc(text)
    for label in label_variants:
        nlabel = nfc(label)
        idx = ntext.find(nlabel)
        if idx != -1:
            return nlabel, idx
    return None, -1


def _build_insertion(label, value):
    if label and label[-1] in ':।ঃ':
        return f" {value}"
    return f": {value}"


def _insert_value(text, label_variants, value, stop_labels=None):
    ntext = nfc(text)
    label, idx = _find_label(ntext, label_variants)
    if label is None:
        return None
    end = idx + len(label)
    remainder = ntext[end:]

    stop_idx = len(remainder)
    found_stop = False
    if stop_labels:
        for stop in stop_labels:
            pos = remainder.find(nfc(stop))
            if pos != -1 and pos < stop_idx:
                stop_idx = pos
                found_stop = True

    insertion = _build_insertion(label, value)
    tail = remainder[stop_idx:]

    if found_stop:
        original_gap_width = stop_idx
        pad = max(original_gap_width - len(insertion), 1)
        insertion = insertion + (" " * pad)
    elif tail and not tail[0].isspace():
        tail = " " + tail

    return ntext[:end] + insertion + tail


def update_text_in_paragraphs(doc, label_variants, new_value):
    for paragraph in doc.paragraphs:
        new_text = _insert_value(paragraph.text, label_variants, new_value)
        if new_text is not None:
            set_paragraph_text_preserving_format(paragraph, new_text)
            return True
    return False


def update_cell_or_paragraph(doc, label_variants, new_value):
    if not new_value:
        return False

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    new_text = _insert_value(paragraph.text, label_variants, new_value)
                    if new_text is not None:
                        set_paragraph_text_preserving_format(paragraph, new_text)
                        return True

    if update_text_in_paragraphs(doc, label_variants, new_value):
        return True

    return False


def fill_header_line(doc, source_no, header_date):
    for paragraph in doc.paragraphs:
        ntext = nfc(paragraph.text)
        if nfc("সুত্র") in ntext and nfc("তারিখ") in ntext:
            text = ntext
            if source_no:
                updated = _insert_value(text, HEADER_SOURCE_NO_LABELS, source_no, stop_labels=HEADER_DATE_LABELS)
                if updated is not None:
                    text = updated

            if header_date:
                updated = _insert_value(text, HEADER_DATE_LABELS, header_date)
                if updated is not None:
                    text = updated

            if text != ntext:
                set_paragraph_text_preserving_format(paragraph, text)
            return True
    return False


def insert_name_in_site_visit_paragraph(doc, institution_name):
    if not institution_name:
        return False

    marker = nfc("এর উপরোক্ত")
    anchor = nfc("সরজমিনে পরিদর্শন")
    replacement = nfc(f"এর {institution_name} উপরোক্ত")

    for paragraph in doc.paragraphs:
        para_text_nfc = nfc(paragraph.text)
        if marker in para_text_nfc and anchor in para_text_nfc:
            new_text = para_text_nfc.replace(marker, replacement, 1)
            set_paragraph_text_preserving_format(paragraph, new_text)
            return True
    return False


def set_cable_quantity(doc, qty):
    if not qty:
        return False

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 4:
                continue
            if "Optical Fiber Cable" in cells[1].text:
                qty_paragraph = cells[3].paragraphs[0]
                set_paragraph_text_preserving_format(qty_paragraph, str(qty))
                return True
    return False


def fill_form(template_path, output_path, data):
    doc = Document(template_path)
    full_address = data.get('name', '')

    fill_map = [
        (LABELS['edc_book'], data.get('source_no', '')),
        (LABELS['connection_nms'], data.get('connection_nms', '')),
        (LABELS['address'], full_address),
        (LABELS['lat_long'], data.get('lat_long', '')),
        (LABELS['union'], data.get('union', '')),
        (LABELS['upazila'], data.get('upazila', '')),
        (LABELS['institution_code'], data.get('institution_code', '')),
        (LABELS['date'], data.get('date', '')),
        (LABELS['router_mac'], data.get('router_mac', '')),
        (LABELS['router_serial'], data.get('router_serial', '')),
    ]

    for label_variants, value in fill_map:
        if value:
            update_cell_or_paragraph(doc, label_variants, value)

    fill_header_line(doc, data.get('source_no', ''), data.get('date', ''))
    insert_name_in_site_visit_paragraph(doc, data.get('name', ''))
    set_cable_quantity(doc, data.get('cable_qty', ''))

    cable_brand = data.get('cable_brand', 'BRB')
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "Optical Fiber Cable Brand:" in paragraph.text:
                        new_text = _insert_value(paragraph.text, ["Optical Fiber Cable Brand:"], cable_brand)
                        if new_text:
                            set_paragraph_text_preserving_format(paragraph, new_text)
                        break

    doc.save(output_path)


def sanitize_filename(name):
    name = re.sub(r'[\\/:*?"<>|]+', '', name)
    name = re.sub(r'\s+', '_', name).strip('_')
    return name[:60] if name else "form"


def _missing_fields_for(data):
    return [(key, label) for key, label in FIELD_PROMPTS if not data.get(key)]


@app.route("/", methods=["GET"])
def index():
    return render_template("index.html")


@app.route("/parse", methods=["POST"])
def parse():
    raw = sanitize_text(request.form.get("raw_text", ""))

    if not raw.strip():
        flash("Paste the contents of your t.txt file first.")
        return redirect(url_for("index"))

    records = split_records(raw)
    if not records:
        flash("No institution records found in the pasted text (looking for blocks separated by a line that just says 'another').")
        return redirect(url_for("index"))

    if not os.path.exists(TEMPLATE_PATH):
        flash(f"Template '{TEMPLATE_PATH}' was not found next to app.py. Add it and try again.")
        return redirect(url_for("index"))

    records_view = []
    for idx, record_text in enumerate(records):
        data = parse_raw_data(record_text)
        known = [(key, label, data.get(key, "")) for key, label in FIELD_PROMPTS if data.get(key)]
        missing = _missing_fields_for(data)
        records_view.append({
            "idx": idx,
            "name": data.get("name") or f"Record {idx + 1}",
            "known": known,
            "missing": missing,
            "always_manual": ALWAYS_MANUAL_FIELDS,
        })

    return render_template("fields.html", records=records_view)


@app.route("/generate", methods=["POST"])
def generate():
    if not os.path.exists(TEMPLATE_PATH):
        flash(f"Template '{TEMPLATE_PATH}' was not found next to app.py.")
        return redirect(url_for("index"))

    try:
        record_count = int(request.form.get("record_count", "0"))
    except ValueError:
        record_count = 0

    if record_count <= 0:
        flash("No records to generate — paste your text again.")
        return redirect(url_for("index"))

    date_mode = request.form.get("date_mode", "auto")
    custom_date_raw = request.form.get("custom_date", "").strip()
    chosen_date = get_today_bangla_date()
    if date_mode == "custom" and custom_date_raw:
        try:
            chosen_dt = datetime.strptime(custom_date_raw, "%Y-%m-%d")
            chosen_date = get_bangla_date(chosen_dt)
        except ValueError:
            chosen_date = get_today_bangla_date()

    output_files = []

    for idx in range(record_count):
        data = {}
        for key, _label in FIELD_PROMPTS:
            data[key] = request.form.get(f"record-{idx}-{key}", "").strip()

        for key, _label, _placeholder, default in ALWAYS_MANUAL_FIELDS:
            val = request.form.get(f"record-{idx}-{key}", "").strip()
            data[key] = val if val else default

        data['date'] = chosen_date

        nms_id = data.get('connection_nms', '').strip()
        if nms_id:
            file_base = sanitize_filename(nms_id)
        else:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            file_base = f"{sanitize_filename(data.get('name', f'record{idx + 1}'))}_{timestamp}"

        output_file = os.path.join(OUTPUT_DIR, f"{file_base}.docx")
        fill_form(TEMPLATE_PATH, output_file, data)
        output_files.append(output_file)

    if len(output_files) == 1:
        return send_file(output_files[0], as_attachment=True,
                          download_name=os.path.basename(output_files[0]))

    mem_zip = io.BytesIO()
    with zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in output_files:
            zf.write(path, arcname=os.path.basename(path))
    mem_zip.seek(0)
    return send_file(mem_zip, as_attachment=True, download_name="edc_forms.zip",
                      mimetype="application/zip")


if __name__ == "__main__":
    app.run(debug=True)
