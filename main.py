import sys
import re
import unicodedata
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

def nfc(s):
    """Normalize Bengali (and any) text to NFC so string matches work
    regardless of how conjuncts like 'য়' were composed in the source.
    IMPORTANT: this must be applied to BOTH sides of every comparison —
    our LABELS constants and text pulled out of the .docx can use
    different (but visually identical) Unicode compositions."""
    return unicodedata.normalize('NFC', s) if s else s

def sanitize_text(s):
    """Strip characters that can't be encoded to UTF-8 — most commonly
    lone/unpaired surrogate code points (U+D800-U+DFFF). These can end
    up in t.txt or in typed input from encoding mismatches (bad copy-
    paste on Windows, a JSON round-trip that preserved an unpaired
    \\uXXXX escape, a mismatched file encoding, etc.). python-docx/lxml
    refuses to serialize them and crashes with
    'UnicodeEncodeError: surrogates not allowed' — this removes the
    offending character(s) instead of failing the whole run."""
    if not s:
        return s
    return s.encode('utf-8', 'ignore').decode('utf-8')

# ============================================================
# 0. FONT HANDLING – Bangla in Nikosh, English/Latin in Times New Roman
# ============================================================
BANGLA_FONT = "Nikosh"
ENGLISH_FONT = "Times New Roman"
DEFAULT_SIZE = Pt(11)  # fallback only; real size is captured from the template

def _is_bangla_char(ch):
    return '\u0980' <= ch <= '\u09FF'

def _char_script(ch):
    """Classify a character as Bangla (True), Latin/other-script (False),
    or script-neutral (None). Bangla digits (০-৯) fall inside the same
    Unicode block as Bangla letters, but Python's str.isalpha() is False
    for digit characters — so a string of ONLY Bangla digits (e.g. an
    EDC Book No. like '১২০০') was never being recognized as Bangla at
    all, since it has no letters to anchor the classification on. This
    checks the Unicode block directly for anything that could plausibly
    be Bangla (letters AND digits AND signs), and only lets truly
    script-neutral characters (spaces, generic punctuation) fall through
    to inherit whatever script surrounds them."""
    if _is_bangla_char(ch):
        return True
    if ch.isalpha() or ch.isdigit():
        return False
    return None

def _split_by_script(text):
    """Split text into consecutive (segment, is_bangla) chunks. Spaces and
    punctuation are attached to whichever segment they fall inside so words
    don't get needlessly fragmented."""
    segments = []
    current = ''
    current_is_bangla = None
    for ch in text:
        is_bangla = _char_script(ch)
        if is_bangla is None:
            is_bangla = current_is_bangla if current_is_bangla is not None else False
        if current_is_bangla is None:
            current_is_bangla = is_bangla
            current = ch
        elif is_bangla == current_is_bangla:
            current += ch
        else:
            segments.append((current, current_is_bangla))
            current = ch
            current_is_bangla = is_bangla
    if current:
        segments.append((current, current_is_bangla))
    return segments if segments else [(text, False)]

def _set_run_font(run, font_name):
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), font_name)

def _set_run_size(run, size):
    """Set BOTH w:sz (regular size) and w:szCs (complex-script size).
    Word renders Bangla/Bengali using w:szCs, not w:sz, since it treats
    Bangla as a 'complex script'. python-docx's run.font.size only ever
    writes w:sz — leaving w:szCs unset, which makes Word fall back to
    the document's default style size for any Bangla text (12pt in this
    template) even though the Latin text correctly showed 11pt."""
    run.font.size = size
    rPr = run._element.get_or_add_rPr()
    half_points = str(int(round(size.pt * 2)))
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), half_points)

def _copy_run_format(src_run, dst_run):
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    _set_run_size(dst_run, src_run.font.size or DEFAULT_SIZE)
    try:
        if src_run.font.color and src_run.font.color.rgb is not None:
            dst_run.font.color.rgb = src_run.font.color.rgb
    except Exception:
        pass

def apply_mixed_fonts(paragraph, run):
    """Given a run whose .text was just set (possibly mixing Bangla and
    English/Latin characters), split it into separate runs so each script
    gets its correct font: Bangla -> Nikosh, English/Latin -> Times New
    Roman. Preserves bold/italic/underline/size/color and run order.
    NOTE: `run` must already carry the correct size/bold/etc (see
    _capture_format / _apply_captured_format below) BEFORE calling this,
    since every split-off segment copies its formatting from `run`."""
    text = run.text
    if not text:
        return
    segments = _split_by_script(text)

    first_text, first_is_bangla = segments[0]
    run.text = first_text
    _set_run_font(run, BANGLA_FONT if first_is_bangla else ENGLISH_FONT)

    prev_element = run._element
    for seg_text, seg_is_bangla in segments[1:]:
        new_run = paragraph.add_run(seg_text)
        _copy_run_format(run, new_run)
        _set_run_font(new_run, BANGLA_FONT if seg_is_bangla else ENGLISH_FONT)
        prev_element.addnext(new_run._element)
        prev_element = new_run._element

# ============================================================
# 0b. FORMAT CAPTURE/RESTORE
# ============================================================
# python-docx's `paragraph.text = "..."` (and `cell.text = "..."`) wipes
# every existing run and replaces it with a single brand-new run that has
# NO explicit formatting — it silently falls back to the style default
# (in this template: Liberation Serif, 12pt), even though every field
# label in the template is directly formatted as Nikosh/Times New Roman,
# 11pt. That mismatch (12pt instead of 11pt) is what was reflowing text
# and breaking table/page layout. Fix: snapshot the real formatting from
# the paragraph's first run BEFORE replacing the text, then re-apply it
# to the new run afterward.
def _capture_format(paragraph):
    if paragraph.runs:
        r = paragraph.runs[0]
        color = None
        try:
            if r.font.color and r.font.color.type is not None:
                color = r.font.color.rgb
        except Exception:
            pass
        return {
            'size': r.font.size or DEFAULT_SIZE,
            'bold': r.font.bold,
            'italic': r.font.italic,
            'underline': r.font.underline,
            'color': color,
        }
    return {'size': DEFAULT_SIZE, 'bold': None, 'italic': None, 'underline': None, 'color': None}

def _apply_captured_format(run, fmt):
    _set_run_size(run, fmt['size'])
    run.font.bold = fmt['bold']
    run.font.italic = fmt['italic']
    run.font.underline = fmt['underline']
    if fmt['color'] is not None:
        try:
            run.font.color.rgb = fmt['color']
        except Exception:
            pass

def set_paragraph_text_preserving_format(paragraph, new_text):
    """Replace a paragraph's text while keeping its original direct
    formatting (size/bold/etc.), then re-split it by script so Bangla
    goes to Nikosh and English/Latin goes to Times New Roman."""
    new_text = sanitize_text(new_text)
    fmt = _capture_format(paragraph)
    paragraph.text = new_text
    if not paragraph.runs:
        return
    run = paragraph.runs[0]
    _apply_captured_format(run, fmt)
    apply_mixed_fonts(paragraph, run)

# ============================================================
# 1. BANGLA DIGIT CONVERTER
# ============================================================
def to_bangla_digits(number):
    bangla_digits = {'0': '০', '1': '১', '2': '২', '3': '৩', '4': '৪',
                     '5': '৫', '6': '৬', '7': '৭', '8': '৮', '9': '৯'}
    return ''.join(bangla_digits.get(ch, ch) for ch in str(number))

def get_today_bangla_date():
    today = datetime.now()
    day = to_bangla_digits(f"{today.day:02d}")
    month = to_bangla_digits(f"{today.month:02d}")
    year = to_bangla_digits(today.year)
    return f"{day}/{month}/{year}"

# ============================================================
# 1b. MAC ADDRESS NORMALIZER
# ============================================================
_MAC_CANDIDATE_RE = re.compile(
    r'(?<![0-9A-Fa-f:\-])'
    r'([0-9A-Fa-f](?:[:\-]?[0-9A-Fa-f]){11})'
    r'(?![0-9A-Fa-f])'
)

def _format_mac(hex_only):
    pairs = [hex_only[i:i + 2] for i in range(0, 12, 2)]
    return ':'.join(pairs).upper()

def normalize_mac(full_text, router_line=""):
    label_match = re.search(r'MAC\s*[:।\-]*\s*([^\n]+)', full_text, re.IGNORECASE)
    if label_match:
        candidate = _MAC_CANDIDATE_RE.search(label_match.group(1))
        if candidate:
            hex_only = re.sub(r'[^0-9A-Fa-f]', '', candidate.group(1))
            if len(hex_only) == 12:
                return _format_mac(hex_only)

    if router_line:
        candidate = _MAC_CANDIDATE_RE.search(router_line)
        if candidate:
            hex_only = re.sub(r'[^0-9A-Fa-f]', '', candidate.group(1))
            if len(hex_only) == 12:
                return _format_mac(hex_only)

    candidate = _MAC_CANDIDATE_RE.search(full_text)
    if candidate:
        hex_only = re.sub(r'[^0-9A-Fa-f]', '', candidate.group(1))
        if len(hex_only) == 12:
            return _format_mac(hex_only)

    return ""

# ============================================================
# 2. LABELS (with variations)
# ============================================================
# source_no / header_date are handled by fill_header_line() instead of the
# generic loop below, because in the template they share ONE paragraph
# ("সুত্র নংঃ  ...spacing...  তারিখঃ") — generic single-label matching
# can't safely fill one without touching the other.
LABELS = {
    "edc_book": ["EDC Book No.:"],
    "connection_nms": ["সংযোগ/NMS আইডিঃ"],
    "address": ["সংযোগ গ্রহণকারী প্রতিষ্টানের নাম ও ঠিকানাঃ"],
    "lat_long": ["Lat & Long:"],
    "union": ["ইউনিয়ন/পৌরসভা (ওয়ার্ড):"],
    "upazila": ["উপজেলা/সিটি কর্পোরেশনঃ"],
    "institution_code": ["প্রতিষ্টানের কোড (EMIS/EIIN/OTHERS):"],
    "date": ["সংযোগ প্রদানের তারিখঃ"],
    "router_mac": ["Wi-Fi Router MAC Address(Upon which API will be called):"],
    "router_serial": ["Wi-Fi Router Serial No.:"],
    "cable_brand": ["Optical Fiber Cable Brand:"]
}

HEADER_SOURCE_NO_LABELS = ["সুত্র নংঃ", "সুত্র নং:", "সুত্র নং"]
HEADER_DATE_LABELS = ["তারিখঃ"]

# ============================================================
# 3. RECORD SPLITTER
# ============================================================
_RECORD_SPLIT_RE = re.compile(r'^\s*another\s*$', re.IGNORECASE | re.MULTILINE)

def split_records(raw_text):
    blocks = _RECORD_SPLIT_RE.split(raw_text)
    return [b.strip('\n') for b in blocks if b.strip()]

# ============================================================
# 4. PARSER
# ============================================================
def parse_raw_data(text):
    data = {}
    text = nfc(text)

    def find_field(pattern, default=""):
        match = re.search(nfc(pattern), text, re.IGNORECASE | re.DOTALL)
        return match.group(1).strip() if match else default

    data['serial'] = find_field(r'প্রতিষ্ঠানের\s*সিরিয়াল\s*নাম্বার\s*[:।\-]*\s*([^\n]+)')
    data['name'] = find_field(r'প্রতিষ্ঠানের\s*নাম\s*[:।\-]*\s*([^\n]+)')
    data['mobile'] = find_field(r'মোবাইল\s*নাম্বার\s*[:।\-]*\s*([^\n]+)')
    data['village'] = find_field(r'গ্রাম\s*[:।\-]*\s*([^\n]+)')
    data['pop'] = find_field(r'পপ\s*এর\s*ঠিকানা\s*[:।\-]*\s*([^\n]+)')
    data['union'] = find_field(r'ইউনিয়নের\s*নাম\s*[:।\-]*\s*([^\n]+)')
    data['upazila'] = find_field(r'থানা\s*[:।\-]*\s*([^\n]+)')

    router_line = find_field(r'রাউটার\s*সিরিয়াল\s*নাম্বার\s*[:।\-]*\s*([^\n]+)')
    serial = ""
    if router_line:
        serial = re.sub(nfc(r'^(S/N|SN|সিরিয়াল)\s*[:।\-]?\s*'), '', router_line, flags=re.IGNORECASE).strip()

    mac = normalize_mac(text, router_line)

    if not serial and router_line:
        if mac:
            stripped = re.sub(r'[0-9A-Fa-f]{2}[:\-]?[0-9A-Fa-f]{2}[:\-]?[0-9A-Fa-f]{2}[:\-]?[0-9A-Fa-f]{2}[:\-]?[0-9A-Fa-f]{2}[:\-]?[0-9A-Fa-f]{2}', '', router_line)
            serial = re.sub(nfc(r'^(S/N|SN|সিরিয়াল)\s*[:।\-]?\s*'), '', stripped, flags=re.IGNORECASE).strip()
        else:
            serial = router_line.strip()

    data['router_serial'] = serial
    data['router_mac'] = mac

    return data

# ============================================================
# 4b. PROMPT FOR MISSING FIELDS
# ============================================================
FIELD_PROMPTS = [
    ('name', "প্রতিষ্ঠানের নাম"),
    ('mobile', "মোবাইল নাম্বার"),
    ('village', "গ্রাম"),
    ('pop', "পপ এর ঠিকানা"),
    ('union', "ইউনিয়নের নাম"),
    ('upazila', "থানা/উপজেলা"),
    ('router_serial', "রাউটার সিরিয়াল নাম্বার"),
    ('router_mac', "MAC address"),
]

def prompt_for_missing_fields(data, record_idx):
    for key, label in FIELD_PROMPTS:
        if not data.get(key):
            val = input(f"  ⚠️ Record {record_idx}: '{label}' not found in t.txt — enter manually (blank to leave empty): ").strip()
            if val:
                data[key] = val
    return data

# ============================================================
# 5. LABEL MATCHING / INSERTION  (NFC-safe, non-destructive)
# ============================================================
def _find_label(text, label_variants):
    """NFC-normalized substring search. Returns (matched_label, index) in
    NFC-normalized `text`, or (None, -1)."""
    ntext = nfc(text)
    for label in label_variants:
        nlabel = nfc(label)
        idx = ntext.find(nlabel)
        if idx != -1:
            return nlabel, idx
    return None, -1

def _build_insertion(label, value):
    # Every label in this template already ends with its own punctuation
    # (":" or the Bangla visarga "ঃ" or "।"), so appending another ":" was
    # producing doubled colons everywhere (e.g. "...called): : XX:XX...").
    if label and label[-1] in ':।ঃ':
        return f" {value}"
    return f": {value}"

def _insert_value(text, label_variants, value):
    """Insert `value` right after whichever label variant is found.
    Only ever consumes a single pre-existing space right after the label
    (to avoid a double space) — never anything further, so it can never
    eat a neighboring field's label or content later in the same line
    (the old bug that deleted 'তারিখঃ' from the header line)."""
    ntext = nfc(text)
    label, idx = _find_label(ntext, label_variants)
    if label is None:
        return None
    end = idx + len(label)
    skip = 1 if ntext[end:end + 1] == ' ' else 0
    insertion = _build_insertion(label, value)
    return ntext[:end] + insertion + ntext[end + skip:]

# ============================================================
# 6. DOCUMENT UPDATER
# ============================================================
def update_text_in_paragraphs(doc, label_variants, new_value):
    for paragraph in doc.paragraphs:
        new_text = _insert_value(paragraph.text, label_variants, new_value)
        if new_text is not None:
            set_paragraph_text_preserving_format(paragraph, new_text)
            print(f"  ✅ Updated '{label_variants[0]}' (in paragraph) → '{new_value}'")
            return True
    return False

def update_cell_or_paragraph(doc, label_variants, new_value):
    if not new_value:
        return False

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    new_text = _insert_value(paragraph.text, label_variants, new_value)
                    if new_text is not None:
                        set_paragraph_text_preserving_format(paragraph, new_text)
                        print(f"  ✅ Updated '{label_variants[0]}' (in table) → '{new_value}'")
                        return True

    if update_text_in_paragraphs(doc, label_variants, new_value):
        return True

    print(f"  ⚠️ Label '{label_variants[0]}' (and alternatives) not found anywhere.")
    return False

def fill_header_line(doc, source_no, header_date):
    """The template has 'সুত্র নংঃ' and 'তারিখঃ' sharing one paragraph,
    separated by wide spacing for visual alignment. Fill both in a single
    pass on that one paragraph so neither insertion can clobber the other."""
    for paragraph in doc.paragraphs:
        ntext = nfc(paragraph.text)
        if nfc("সুত্র") in ntext and nfc("তারিখ") in ntext:
            text = ntext
            if source_no:
                updated = _insert_value(text, HEADER_SOURCE_NO_LABELS, source_no)
                if updated is not None:
                    text = updated
                else:
                    print("  ⚠️ 'সুত্র নংঃ' not found in header line.")
            else:
                print("  ⏭️ Skipping 'সুত্র নংঃ' (no value)")

            if header_date:
                updated = _insert_value(text, HEADER_DATE_LABELS, header_date)
                if updated is not None:
                    text = updated
                else:
                    print("  ⚠️ 'তারিখঃ' not found in header line.")
            else:
                print("  ⏭️ Skipping header 'তারিখঃ' (no value)")

            if text != ntext:
                set_paragraph_text_preserving_format(paragraph, text)
                print(f"  ✅ Updated header line → সুত্র নং='{source_no}', তারিখ='{header_date}'")
            return True

    print("  ⚠️ Header line ('সুত্র নংঃ ... তারিখঃ') not found.")
    return False

def insert_name_in_site_visit_paragraph(doc, institution_name):
    if not institution_name:
        print("  ⏭️ Skipping site-visit paragraph insert (no institution name).")
        return False

    marker = nfc("এর উপরোক্ত")
    anchor = nfc("সরজমিনে পরিদর্শন")
    replacement = nfc(f"এর {institution_name} উপরোক্ত")

    for paragraph in doc.paragraphs:
        para_text_nfc = nfc(paragraph.text)
        if marker in para_text_nfc and anchor in para_text_nfc:
            new_text = para_text_nfc.replace(marker, replacement, 1)
            set_paragraph_text_preserving_format(paragraph, new_text)
            print(f"  ✅ Inserted institution name into site-visit paragraph → '{institution_name}'")
            return True

    print("  ⚠️ Site-visit paragraph not found — institution name not inserted.")
    return False

def set_cable_quantity(doc, qty):
    """The equipment table's 'Optical Fiber Cable ... (mtr.)' row has a
    Qty. column hardcoded to 1500 in the template. This finds that row
    by its description text ('Optical Fiber Cable') and overwrites just
    the Qty cell, leaving the description/UoM/Remarks cells untouched."""
    if not qty:
        print("  ⏭️ Skipping Optical Fiber Cable Qty. (no value)")
        return False

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 4:
                continue
            if "Optical Fiber Cable" in cells[1].text:
                qty_paragraph = cells[3].paragraphs[0]
                set_paragraph_text_preserving_format(qty_paragraph, str(qty))
                print(f"  ✅ Updated Optical Fiber Cable Qty. → '{qty}'")
                return True

    print("  ⚠️ Optical Fiber Cable row not found — Qty. not updated.")
    return False

def fill_form(template_path, output_path, data):
    doc = Document(template_path)
    full_address = data.get('name', '')

    fill_map = [
        (LABELS['edc_book'], data.get('source_no', '')),
        (LABELS['connection_nms'], data.get('connection_nms', '')),
        (LABELS['address'], full_address),
        (LABELS['lat_long'], data.get('lat_long', '')),
        (LABELS['union'], data.get('union', '')),
        (LABELS['upazila'], data.get('upazila', '')),
        (LABELS['institution_code'], data.get('institution_code', '')),
        (LABELS['date'], data.get('date', '')),
        (LABELS['router_mac'], data.get('router_mac', '')),
        (LABELS['router_serial'], data.get('router_serial', '')),
    ]

    print("\n🔍 Filling fields:")
    for label_variants, value in fill_map:
        if value:
            update_cell_or_paragraph(doc, label_variants, value)
        else:
            print(f"  ⏭️ Skipping '{label_variants[0]}' (no value)")

    fill_header_line(doc, data.get('source_no', ''), data.get('date', ''))
    insert_name_in_site_visit_paragraph(doc, data.get('name', ''))
    set_cable_quantity(doc, data.get('cable_qty', ''))

    cable_brand = data.get('cable_brand', 'BRB')
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "Optical Fiber Cable Brand:" in paragraph.text:
                        new_text = _insert_value(paragraph.text, ["Optical Fiber Cable Brand:"], cable_brand)
                        if new_text:
                            set_paragraph_text_preserving_format(paragraph, new_text)
                            print(f"  ✅ Updated Optical Fiber Cable Brand → '{cable_brand}'")
                        break

    doc.save(output_path)
    print(f"\n✅ Form saved as '{output_path}'")

# ============================================================
# 7. MAIN
# ============================================================
def sanitize_filename(name):
    name = re.sub(r'[\\/:*?"<>|]+', '', name)
    name = re.sub(r'\s+', '_', name).strip('_')
    return name[:60] if name else "form"

def main():
    try:
        with open("t.txt", "r", encoding="utf-8") as f:
            raw = f.read()
    except FileNotFoundError:
        print("❌ File 't.txt' not found.")
        sys.exit(1)

    raw = sanitize_text(raw)

    if not raw.strip():
        print("❌ 't.txt' is empty.")
        sys.exit(1)

    records = split_records(raw)
    if not records:
        print("❌ No institution records found in 't.txt'.")
        sys.exit(1)

    print(f"📄 Found {len(records)} institution record(s) in t.txt.")

    template = "EDC Form.docx"
    if not os.path.exists(template):
        print(f"❌ Template '{template}' not found.")
        sys.exit(1)

    for idx, record_text in enumerate(records, start=1):
        data = parse_raw_data(record_text)

        print(f"\n{'='*60}")
        print(f"📝 Record {idx}/{len(records)}: {data.get('name') or '(name not found)'}")
        print(f"{'='*60}")

        data = prompt_for_missing_fields(data, idx)

        edc = input("  EDC Book No. / সুত্র নং (e.g., 2026-001): ").strip()
        if edc:
            data['source_no'] = edc

        nms = input("  সংযোগ/NMS আইডিঃ: ").strip()
        if nms:
            data['connection_nms'] = nms

        inst_code = input("  প্রতিষ্টানের কোড (EMIS/EIIN/OTHERS): ").strip()
        if inst_code:
            data['institution_code'] = inst_code

        cable = input("  Optical Fiber Cable Brand (default 'BRB'): ").strip()
        data['cable_brand'] = cable if cable else data.get('cable_brand', 'BRB')

        cable_qty = input("  Optical Fiber Cable Qty. in meters (default '1500'): ").strip()
        data['cable_qty'] = cable_qty if cable_qty else "1500"

        lat = input("  Lat & Long (if available): ").strip()
        if lat:
            data['lat_long'] = lat

        data['date'] = get_today_bangla_date()

        print("\n📦 Final data:")
        for key, val in data.items():
            if val:
                print(f"  {key}: {val}")

        nms_id = data.get('connection_nms', '').strip()
        if nms_id:
            file_base = sanitize_filename(nms_id)
        else:
            print("  ⚠️ No NMS ID entered — falling back to institution name for filename.")
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            file_base = f"{sanitize_filename(data.get('name', f'record{idx}'))}_{timestamp}"
        output_file = f"{file_base}.docx"
        print(f"\n💾 Will save as: {output_file}")

        fill_form(template, output_file, data)

if __name__ == "__main__":
    main()
